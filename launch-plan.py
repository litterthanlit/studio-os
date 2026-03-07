"""
Studio OS — Launch Plan & Waitlist Growth Playbook
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

OUTPUT = "/sessions/nifty-relaxed-wozniak/mnt/studio OS/Studio OS — Launch Plan.docx"

doc = Document()

# ── Page setup ───────────────────────────────────────────────────────────────
section = doc.sections[0]
section.page_width  = Inches(8.5)
section.page_height = Inches(11)
section.left_margin = section.right_margin = Inches(1.1)
section.top_margin  = section.bottom_margin = Inches(1.0)

# ── Colour palette ───────────────────────────────────────────────────────────
NAVY   = RGBColor(0x24, 0x30, 0xAD)
BLACK  = RGBColor(0x11, 0x11, 0x11)
DARK   = RGBColor(0x33, 0x33, 0x33)
MID    = RGBColor(0x55, 0x55, 0x55)
LIGHT  = RGBColor(0x88, 0x88, 0x88)
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
ACCENT = RGBColor(0x81, 0x8C, 0xF8)  # accessible indigo

# ── Helper functions ─────────────────────────────────────────────────────────
def set_para_spacing(para, before=0, after=6, line=None):
    pf = para.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after  = Pt(after)
    if line:
        from docx.shared import Pt as _Pt
        pf.line_spacing = _Pt(line)

def add_run(para, text, bold=False, italic=False, size=11, color=None, font="Calibri"):
    run = para.add_run(text)
    run.bold   = bold
    run.italic = italic
    run.font.name = font
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    return run

def heading1(doc, text):
    p = doc.add_paragraph()
    set_para_spacing(p, before=18, after=6)
    run = p.add_run(text)
    run.bold = True
    run.font.name = "Calibri"
    run.font.size = Pt(18)
    run.font.color.rgb = BLACK
    # bottom border
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '4')
    bottom.set(qn('w:color'), '2430AD')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p

def heading2(doc, text):
    p = doc.add_paragraph()
    set_para_spacing(p, before=12, after=4)
    run = p.add_run(text)
    run.bold = True
    run.font.name = "Calibri"
    run.font.size = Pt(13)
    run.font.color.rgb = NAVY
    return p

def body(doc, text, color=None):
    p = doc.add_paragraph()
    set_para_spacing(p, before=0, after=5, line=13)
    add_run(p, text, size=11, color=color or DARK)
    return p

def bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    set_para_spacing(p, before=0, after=3)
    pf = p.paragraph_format
    pf.left_indent = Inches(0.3)
    if bold_prefix:
        add_run(p, bold_prefix + " ", bold=True, size=11, color=BLACK)
        add_run(p, text, size=11, color=DARK)
    else:
        add_run(p, text, size=11, color=DARK)
    return p

def label(doc, text):
    p = doc.add_paragraph()
    set_para_spacing(p, before=8, after=2)
    run = p.add_run(text.upper())
    run.bold = True
    run.font.name = "Calibri"
    run.font.size = Pt(8)
    run.font.color.rgb = LIGHT
    run.font.all_caps = True
    p.paragraph_format.left_indent = Pt(0)
    return p

def callout(doc, text, bg_hex="F0F4FF"):
    """Shaded callout box."""
    p = doc.add_paragraph()
    set_para_spacing(p, before=8, after=8)
    add_run(p, text, size=10.5, color=DARK)
    pPr = p._p.get_or_add_pPr()
    shading = OxmlElement('w:shd')
    shading.set(qn('w:val'), 'clear')
    shading.set(qn('w:color'), 'auto')
    shading.set(qn('w:fill'), bg_hex)
    pPr.append(shading)
    pf = p.paragraph_format
    pf.left_indent = pf.right_indent = Inches(0.2)
    return p

def divider(doc):
    p = doc.add_paragraph()
    set_para_spacing(p, before=6, after=6)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '2')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'E0E0E0')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p

def page_break(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(docx.oxml.ns.qn and None or None)
    p.clear()
    p.add_run().add_break(docx.enum.text.WD_BREAK_TYPE if False else None)
    # simpler:
    from docx.enum.text import WD_BREAK_TYPE
    p2 = doc.add_paragraph()
    run2 = p2.add_run()
    from docx.oxml import OxmlElement as OE
    br = OE('w:br')
    br.set(qn('w:type'), 'page')
    run2._r.append(br)
    return p2

# ══════════════════════════════════════════════════════════════════════════════
# COVER PAGE
# ══════════════════════════════════════════════════════════════════════════════

p = doc.add_paragraph()
set_para_spacing(p, before=60, after=0)
run = p.add_run("Studio OS")
run.bold = True
run.font.name = "Calibri"
run.font.size = Pt(32)
run.font.color.rgb = BLACK

p2 = doc.add_paragraph()
set_para_spacing(p2, before=4, after=4)
add_run(p2, "Launch Plan & Waitlist Growth Playbook", size=16, color=NAVY, bold=False)

p3 = doc.add_paragraph()
set_para_spacing(p3, before=0, after=40)
add_run(p3, "February 2026  ·  Confidential", size=10, color=LIGHT)

divider(doc)

callout(doc,
    "This document covers: product positioning · MCP strategy & costs · "
    "go-to-market phases · waitlist growth tactics · path to £/$300K ARR.",
    bg_hex="F0F4FF")

# ══════════════════════════════════════════════════════════════════════════════
# 1. SITUATION & HONEST ASSESSMENT
# ══════════════════════════════════════════════════════════════════════════════

heading1(doc, "1.  Situation & Honest Assessment")

body(doc,
    "Studio OS is an AI-powered design workspace: it pulls references from Pinterest, "
    "Are.na, and Lummi; scores them on composition, colour, and mood; builds design "
    "system tokens; and exports AI-ready briefs. The product is functionally ahead of "
    "the audience's awareness of it.")

heading2(doc, "What you have")
for item in [
    ("2 completed case studies", "VCEEZY Framer site, Figma course work."),
    ("2 templates in development", "Flowline (SaaS), Terrene (Architecture)."),
    ("Studio OS app", "Home, sidebar, command palette, animation system, AI curation — all built."),
    ("Framer expertise", "Rare skill. High demand. Defensible."),
    ("Revenue ambition", "$300K ARR target for 2026."),
]:
    bullet(doc, item[1], bold_prefix=item[0])

heading2(doc, "What's missing right now")
for item in [
    "No waitlist — the CTA exists but zero collection infrastructure in place",
    "No public URL — the product is invisible to the market",
    "No validated audience — 'solo founders, startups, art galleries' needs a test",
    "No consistent output — goal is 1 template/week, current pace is 2 WIP",
    "No case studies 3–8 — need 6 more to hit the social proof target",
]:
    bullet(doc, item)

callout(doc,
    "The risk: building a Mercedes when the market wants a Honda. The fix: "
    "deploy, put the waitlist in front of real designers, and let signal tell you "
    "which features to double down on. Revenue follows distribution — not polish.",
    bg_hex="FFF8F0")

# ══════════════════════════════════════════════════════════════════════════════
# 2. POSITIONING
# ══════════════════════════════════════════════════════════════════════════════

heading1(doc, "2.  Positioning")

body(doc,
    "Positioning is a claim you make and then prove with the product. "
    "Good positioning is exclusive — if your competitor could say the same thing, rewrite it.")

heading2(doc, "Recommended position")
callout(doc,
    "Studio OS is the first design workspace built around AI curation — "
    "not task management, not file storage. It scores your references, "
    "builds your system, and ships your brief. For designers who work with AI tools daily.",
    bg_hex="F0F4FF")

heading2(doc, "Who it's for (primary)")
for row in [
    ("Freelance designers", "Solo practitioners juggling client work. High pain around brief generation and handing off to devs. Willing to pay for time saved."),
    ("Small design studios (2–8 people)", "Need a shared reference and token layer. Currently cobbling together Notion + Figma + Pinterest. Budget exists."),
    ("Designers using AI daily", "Cursor, Midjourney, Claude, Framer AI users. Already sold on AI-assisted workflows. Studio OS is infrastructure for them."),
]:
    bullet(doc, row[1], bold_prefix=row[0])

heading2(doc, "Who it's NOT for (yet)")
for item in [
    "Enterprise design teams — too much complexity, procurement cycles too long",
    "Non-designers — the product assumes design literacy",
    "Art galleries — interesting long-term, too niche for launch",
]:
    bullet(doc, item)

# ══════════════════════════════════════════════════════════════════════════════
# 3. MCP STRATEGY & COSTS
# ══════════════════════════════════════════════════════════════════════════════

heading1(doc, "3.  MCP Strategy & Costs")

body(doc,
    "MCP (Model Context Protocol) is Anthropic's open standard for giving AI tools "
    "structured access to external data and actions. Building an MCP server for "
    "Studio OS means any MCP-compatible AI client — Claude Desktop, Cursor, Windsurf, "
    "Zed, and others — can read your projects, tokens, and briefs directly.")

heading2(doc, "What an MCP server unlocks for Studio OS")
for item in [
    "Claude Desktop reads your active project tokens while you write copy or generate images",
    "Cursor pulls your design-system.md as live context during frontend development",
    "Future: any AI tool can write back to Studio OS (log a component, update a token)",
    "Massive distribution signal — every MCP directory listing is free reach to AI-native developers",
]:
    bullet(doc, item)

heading2(doc, "Build cost: effectively zero")
body(doc,
    "The MCP SDK (Python or TypeScript) is free and open-source. "
    "Building the server is a 1–2 day engineering task — expose 4–6 tools "
    "(list_projects, get_tokens, get_brief, add_reference) and you have a "
    "fully functional MCP server. No infrastructure cost beyond what you already run.")

heading2(doc, "Claude API costs (the real number)")

# Table
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
table = doc.add_table(rows=1, cols=4)
table.style = 'Table Grid'
hdr = table.rows[0].cells
for i, txt in enumerate(["Feature", "Model", "Est. tokens/call", "Cost/1k calls"]):
    hdr[i].text = txt
    run = hdr[i].paragraphs[0].runs[0]
    run.bold = True
    run.font.size = Pt(9.5)
    run.font.color.rgb = WHITE
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), '2430AD')
    hdr[i].paragraphs[0]._p.get_or_add_pPr().append(shd)

rows_data = [
    ("Brief generation", "Sonnet 4.5", "3k in + 1k out", "~$24"),
    ("Image scoring (batch 10)", "Haiku 4.5", "2k in + 0.5k out", "~$1.60"),
    ("Token suggestions", "Haiku 4.5", "1k in + 0.5k out", "~$0.80"),
    ("MCP query responses", "Haiku 4.5", "1.5k in + 0.3k out", "~$1.00"),
]
for row_data in rows_data:
    row = table.add_row().cells
    for i, val in enumerate(row_data):
        row[i].text = val
        row[i].paragraphs[0].runs[0].font.size = Pt(9.5)

p_after = doc.add_paragraph()
set_para_spacing(p_after, before=4, after=4)

body(doc,
    "At 500 monthly active users each generating 2 briefs and scoring 50 images/month: "
    "brief cost ~$24, image scoring ~$8 — total API cost under $35/month. "
    "Margins stay healthy even at $19/month pricing. At $49/month the unit economics are excellent.")

callout(doc,
    "Recommendation: build the MCP server in Phase 2 (post-launch). "
    "Lead with it as a distribution play — list on glama.ai, mcp.so, and the Anthropic MCP directory. "
    "It will drive inbound from AI-native developers who become your most vocal early users.",
    bg_hex="F0F4FF")

# ══════════════════════════════════════════════════════════════════════════════
# 4. GO-TO-MARKET PHASES
# ══════════════════════════════════════════════════════════════════════════════

heading1(doc, "4.  Go-to-Market Phases")

phases = [
    ("Phase 1 — Foundation", "Now → 4 weeks",
     "F8F9FF", [
        "Deploy Studio OS to Vercel (required for Pinterest OAuth + Lummi rate limit bump)",
        "Set up waitlist API (done) — connect to Resend audience for email sequences",
        "Ship Flowline template — first public, paid Framer asset",
        "Post 3 case studies to portfolio (VCEEZY + 2 new)",
        "Create @studioOS Twitter/X account and post first 5 'build in public' threads",
    ]),
    ("Phase 2 — Audience", "Weeks 5–10",
     "F0FFF4", [
        "Weekly Framer template drop (Terrene + 2 more) — build catalogue",
        "Build MCP server — list on glama.ai, mcp.so, Anthropic directory",
        "Launch on Product Hunt (prep 3 weeks ahead, brief designer community)",
        "Post 1 design system breakdown/week on Twitter — become the 'design tokens for AI' account",
        "Start Are.na channel: curated weekly brief from Studio OS AI scoring",
        "Target: 500 waitlist sign-ups before full launch",
    ]),
    ("Phase 3 — Revenue", "Weeks 11–20",
     "FFF8F0", [
        "Open Studio OS to waitlist — $0/month early access, then $19/month",
        "Ship 2 premium Framer templates/month at $79–$149 each",
        "Offer 'design system setup' service ($500–$1,500 per client) to generate cash while SaaS scales",
        "Affiliate programme: 30% recurring for template/plugin referrals",
        "Target: $5K MRR by end of Phase 3",
    ]),
    ("Phase 4 — Scale", "Months 6–12",
     "FFF0F8", [
        "Studio OS Pro tier: team workspaces, shared token libraries, client portal",
        "Framer partner listing — official visibility in Framer marketplace",
        "Agency bundle: white-label briefs, client handoff mode",
        "B2B outbound to small studios (2–5 person teams) via LinkedIn + cold email",
        "Target: $25K MRR → run-rate of $300K ARR",
    ]),
]

for phase_name, timeline, color, items in phases:
    heading2(doc, f"{phase_name}  ·  {timeline}")
    for item in items:
        bullet(doc, item)
    doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# 5. WAITLIST GROWTH PLAYBOOK
# ══════════════════════════════════════════════════════════════════════════════

heading1(doc, "5.  Waitlist Growth Playbook")

body(doc,
    "The goal is 500 qualified sign-ups before opening access — 'qualified' means "
    "designers who use AI tools and have heard of Are.na or Framer. Quality over quantity.")

heading2(doc, "Channel 1: Twitter/X — Build in Public")
body(doc, "Your highest-leverage channel. Designers are on Twitter. AI builders are on Twitter. The overlap is your audience.")
for item in [
    "Post the Studio OS homepage — one screenshot, one sentence. No hype.",
    "Weekly 'design system breakdown' thread: take a well-known brand (Notion, Linear, Arc) and reverse-engineer their tokens. End every thread with a Studio OS waitlist link.",
    "Post AI scoring in action — screenshot of your inspiration feed with the scores visible. Caption: 'This is how I curate references now.'",
    "Reply to every @framer, @linearapp, @vercel thread with genuine insight — not promotion. Build the account before you sell from it.",
    "Target accounts to engage: @rauchg, @delba_oliveira, @jsngr, @tinyxr, @frankiecd86",
]:
    bullet(doc, item)

heading2(doc, "Channel 2: Are.na — Credibility play")
body(doc, "Are.na has a small but extremely influential design audience. Studio OS integrates with it — that's a natural hook.")
for item in [
    "Create a public Studio OS channel on Are.na: 'AI-curated references, scored weekly'",
    "Every Friday: post your top 5 images from the week's Studio OS scoring with brief commentary",
    "DM Are.na's most-followed curators — offer early access in exchange for honest feedback",
]:
    bullet(doc, item)

heading2(doc, "Channel 3: Framer Community")
for item in [
    "Post Flowline and Terrene in Framer's official template marketplace",
    "Join Framer Discord — answer questions about design systems and tokens, mention Studio OS naturally",
    "Post a 'how I built this template with Studio OS' Loom video — no script, just screen-share",
    "Submit a guest post to Framer's blog: 'Building a design system-first Framer template'",
]:
    bullet(doc, item)

heading2(doc, "Channel 4: Product Hunt prep")
body(doc, "Do not launch cold. Prep 3 weeks ahead:")
for item in [
    "Build a 'coming soon' PH page 3 weeks before launch date",
    "DM 20 designers/indie hackers you admire — ask them to upvote on launch day",
    "Post the launch on Twitter the morning it goes live — include a GIF of the product",
    "Aim for top 5 Product of the Day — this alone can generate 200–400 sign-ups",
]:
    bullet(doc, item)

heading2(doc, "Channel 5: Cold outreach (small studios)")
body(doc, "High effort, high conversion. Target 2–5 person design studios who post work on Dribbble or Instagram.")
for item in [
    "Find studios using Framer (check their site source) — they're pre-qualified",
    "Send a personalised 3-line DM: 'Saw your [project]. I'm building a workspace that does X. Happy to give you early access if useful.'",
    "50 DMs/week → expect 5–8 replies → 2–3 sign-ups. Low volume but high signal.",
]:
    bullet(doc, item)

heading2(doc, "Waitlist email sequence (post sign-up)")
for i, (subject, body_text) in enumerate([
    ("You're on the Studio OS list", "Thanks for joining. Here's what Studio OS actually does [30-second GIF]. We'll be in touch when early access opens."),
    ("How we score your references (and why it matters)", "Behind the scenes on the AI scoring system. Composition, colour, mood — and how the threshold of 80 was chosen."),
    ("Building your first design system in Studio OS", "Walkthrough of the token builder. 3-minute Loom. No fluff."),
    ("Early access is opening — here's your link", "You're in. Here's how to get started in 5 minutes."),
], 1):
    bullet(doc, f"Email {i}: {subject} — {body_text}")

# ══════════════════════════════════════════════════════════════════════════════
# 6. REVENUE PATH TO $300K
# ══════════════════════════════════════════════════════════════════════════════

heading1(doc, "6.  Revenue Path to $300K ARR")

table2 = doc.add_table(rows=1, cols=4)
table2.style = 'Table Grid'
hdr2 = table2.rows[0].cells
for i, txt in enumerate(["Stream", "Price", "Volume needed", "Annual revenue"]):
    hdr2[i].text = txt
    run = hdr2[i].paragraphs[0].runs[0]
    run.bold = True
    run.font.size = Pt(9.5)
    run.font.color.rgb = WHITE
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), '2430AD')
    hdr2[i].paragraphs[0]._p.get_or_add_pPr().append(shd)

revenue_rows = [
    ("Studio OS SaaS", "$19/mo", "600 users", "$136,800"),
    ("Studio OS Pro", "$49/mo", "150 users", "$88,200"),
    ("Framer templates", "$99 avg", "300 sales", "$29,700"),
    ("Design system services", "$750 avg", "60 projects", "$45,000"),
    ("Affiliate/referrals", "–", "–", "~$10,000"),
    ("TOTAL", "", "", "$309,700"),
]
for row_data in revenue_rows:
    row = table2.add_row().cells
    for i, val in enumerate(row_data):
        row[i].text = val
        run = row[i].paragraphs[0].runs[0] if row[i].paragraphs[0].runs else row[i].paragraphs[0].add_run(val)
        run.font.size = Pt(9.5)
        if row_data[0] == "TOTAL":
            run.bold = True

p_after2 = doc.add_paragraph()
set_para_spacing(p_after2, before=4, after=4)

callout(doc,
    "The services line is the fastest path to early revenue. "
    "1 design system setup per week at $750 = $39K/year and gives you direct "
    "user feedback. Do this in Phase 2 while SaaS scales.",
    bg_hex="F0FFF4")

# ══════════════════════════════════════════════════════════════════════════════
# 7. IMMEDIATE NEXT ACTIONS
# ══════════════════════════════════════════════════════════════════════════════

heading1(doc, "7.  Immediate Next Actions (This Week)")

for i, item in enumerate([
    "Deploy to Vercel — the app needs a live URL for OAuth and distribution",
    "Ship the waitlist API to production — emails are currently going nowhere",
    "Create Twitter/X account — post the homepage screenshot today",
    "Finish Flowline template — first paid product, creates income signal",
    "Write case study #3 — one new one per week until you have 8",
    "DM 10 designers on Are.na — offer early access, ask for a 15-minute call",
], 1):
    bullet(doc, item, bold_prefix=f"{i}.")

divider(doc)

p_end = doc.add_paragraph()
set_para_spacing(p_end, before=10, after=0)
add_run(p_end, "The product is ready. The question is distribution. Start there.",
        size=12, bold=True, color=BLACK)

# ── Save ─────────────────────────────────────────────────────────────────────
doc.save(OUTPUT)
print(f"Saved: {OUTPUT}")
